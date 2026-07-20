"""
Module de gestion des fichiers uploadés
Supporte : PDF, Word (.docx), Images, Code, Markdown, Texte
"""

import io
import base64
import importlib
from pathlib import Path
from typing import Dict

# ── Extraction PDF ──────────────────────────────────────────────
def extraire_pdf(contenu_bytes: bytes) -> str:
    """Extrait le texte d'un PDF"""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(contenu_bytes)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages, 1):
                texte = page.extract_text()
                if texte and texte.strip():
                    pages.append(f"[Page {i}]\n{texte.strip()}")
            return "\n\n".join(pages) if pages else "PDF vide ou non lisible."
    except ImportError:
        try:
            pypdf = importlib.import_module("pypdf")
            reader = pypdf.PdfReader(io.BytesIO(contenu_bytes))
            pages = []
            for i, page in enumerate(reader.pages, 1):
                texte = page.extract_text()
                if texte:
                    pages.append(f"[Page {i}]\n{texte.strip()}")
            return "\n\n".join(pages) if pages else "PDF vide."
        except Exception as e:
            return f"Erreur extraction PDF : {e}"
    except Exception as e:
        return f"Erreur PDF : {e}"


# ── Extraction Word (.docx) ─────────────────────────────────────
def extraire_docx(contenu_bytes: bytes) -> str:
    """Extrait le texte d'un document Word (.docx)"""
    try:
        import docx  # python-docx
    except ImportError:
        return ("[Le module python-docx n'est pas installé. "
                "Lance : pip install python-docx]")
    try:
        document = docx.Document(io.BytesIO(contenu_bytes))
        blocs = []

        # Paragraphes
        for para in document.paragraphs:
            if para.text and para.text.strip():
                blocs.append(para.text.strip())

        # Tableaux
        for table in document.tables:
            for ligne in table.rows:
                cellules = [c.text.strip() for c in ligne.cells if c.text.strip()]
                if cellules:
                    blocs.append(" | ".join(cellules))

        texte = "\n".join(blocs)
        return texte if texte.strip() else "Document Word vide ou sans texte extractible."
    except Exception as e:
        return f"Erreur extraction Word : {e}"


# ── Extraction Image ────────────────────────────────────────────
def encoder_image_base64(contenu_bytes: bytes, media_type: str) -> str:
    """Encode une image en base64 pour envoi au LLM"""
    return base64.b64encode(contenu_bytes).decode("utf-8")


# ── Extraction Code & Texte ─────────────────────────────────────
def extraire_texte(contenu_bytes: bytes, nom_fichier: str) -> str:
    """Extrait le contenu texte/code d'un fichier"""
    try:
        return contenu_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return contenu_bytes.decode("latin-1")
        except Exception:
            return f"Impossible de lire le fichier {nom_fichier}"


# ── Détection type de fichier ───────────────────────────────────
EXTENSIONS_CODE = {
    ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".java": "java", ".c": "c", ".cpp": "c++", ".h": "c",
    ".go": "go", ".rs": "rust", ".php": "php", ".rb": "ruby",
    ".cs": "csharp", ".kt": "kotlin", ".swift": "swift",
    ".sh": "bash", ".sql": "sql", ".r": "r", ".dart": "dart",
}

EXTENSIONS_DOC = {
    ".md": "markdown", ".txt": "texte", ".rst": "rst",
    ".yaml": "yaml", ".yml": "yaml", ".json": "json",
    ".toml": "toml", ".xml": "xml", ".html": "html", ".css": "css",
}

EXTENSIONS_WORD  = {".docx", ".doc"}
EXTENSIONS_IMAGE = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}
EXTENSIONS_PDF   = {".pdf"}
EXTENSIONS_AUDIO = {".wav", ".mp3", ".m4a", ".ogg", ".webm", ".flac", ".mpeg", ".mp4"}


def detecter_type(nom_fichier: str) -> str:
    """Retourne le type du fichier"""
    ext = Path(nom_fichier).suffix.lower()
    if ext in EXTENSIONS_PDF:    return "pdf"
    if ext in EXTENSIONS_WORD:   return "word"
    if ext in EXTENSIONS_IMAGE:  return "image"
    if ext in EXTENSIONS_AUDIO:  return "audio"
    if ext in EXTENSIONS_CODE:   return "code"
    if ext in EXTENSIONS_DOC:    return "doc"
    return "inconnu"


def get_langage(nom_fichier: str) -> str:
    """Retourne le langage de programmation"""
    ext = Path(nom_fichier).suffix.lower()
    return EXTENSIONS_CODE.get(ext, EXTENSIONS_DOC.get(ext, "texte"))


# ── Traitement principal ────────────────────────────────────────
def traiter_fichier(fichier_uploade) -> Dict:
    """
    Traite un fichier uploadé depuis Streamlit.
    Retourne un dict avec type, contenu, métadonnées.
    Applique les limites de taille / extensions (sécurité).
    """
    import sys
    from pathlib import Path as _P
    sys.path.insert(0, str(_P(__file__).resolve().parent.parent))
    import config as _cfg

    nom     = fichier_uploade.name
    ext     = Path(nom).suffix.lower()
    contenu = fichier_uploade.read()
    taille  = len(contenu)
    type_f  = detecter_type(nom)
    langage = get_langage(nom)

    resultat = {
        "nom"    : nom,
        "type"   : type_f,
        "langage": langage,
        "taille" : taille,
        "contenu": "",
        "image_b64": None,
        "media_type": None,
        "erreur" : None,
    }

    if ext and ext not in _cfg.UPLOAD_EXTENSIONS_OK:
        resultat["erreur"] = f"Extension non autorisée : {ext}"
        resultat["contenu"] = resultat["erreur"]
        return resultat

    if taille > _cfg.UPLOAD_MAX_BYTES:
        mo = _cfg.UPLOAD_MAX_BYTES / (1024 * 1024)
        resultat["erreur"] = f"Fichier trop volumineux (max {mo:.0f} Mo)."
        resultat["contenu"] = resultat["erreur"]
        return resultat

    try:
        if type_f == "pdf":
            resultat["contenu"] = extraire_pdf(contenu)

        elif type_f == "word":
            texte = extraire_docx(contenu)
            if len(texte) > 8000:
                texte = texte[:8000] + f"\n\n[... document tronqué à 8000 caractères sur {len(texte)} ...]"
            resultat["contenu"] = texte

        elif type_f == "image":
            ext = Path(nom).suffix.lower()
            media_map = {
                ".png": "image/png", ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg", ".gif": "image/gif",
                ".webp": "image/webp",
            }
            resultat["media_type"] = media_map.get(ext, "image/png")
            resultat["image_b64"]  = encoder_image_base64(contenu, resultat["media_type"])
            resultat["contenu"]    = f"[Image : {nom}]"

        elif type_f == "audio":
            resultat["contenu"] = f"[Audio : {nom}]"
            resultat["audio_bytes"] = contenu

        elif type_f in ("code", "doc"):
            texte = extraire_texte(contenu, nom)
            if len(texte) > 8000:
                texte = texte[:8000] + f"\n\n[... fichier tronqué à 8000 caractères sur {len(texte)} ...]"
            resultat["contenu"] = texte

        else:
            # Type inconnu : on tente une extraction texte, mais on prévient
            texte = extraire_texte(contenu, nom)
            # Si le texte décodé contient trop de caractères binaires, on signale
            if texte.count("\ufffd") > 20 or not texte.strip():
                resultat["contenu"] = (
                    f"[Le format du fichier '{nom}' n'est pas pris en charge "
                    f"pour l'extraction de texte. Formats supportés : PDF, Word (.docx), "
                    f"code, Markdown, texte.]"
                )
            else:
                resultat["contenu"] = texte

    except Exception as e:
        resultat["erreur"]  = str(e)
        resultat["contenu"] = f"Erreur lors du traitement : {e}"

    return resultat


def analyser_images_vision(
    fichiers_traites: list,
    question: str = "",
    llm_client=None,
) -> list:
    """
    Remplace le placeholder image par une vraie analyse multimodale Groq.
    Enrichit chaque dict fichier (contenu + vision_ok).
    """
    if not fichiers_traites:
        return fichiers_traites
    client = llm_client
    if client is None:
        try:
            from generation.llm_client import LLMClient

            client = LLMClient()
        except Exception:
            return fichiers_traites

    q_user = (question or "").strip()
    if q_user:
        prompt_vision = (
            "L'utilisateur a joint cette image et demande :\n"
            f"{q_user}\n\n"
            "Réponds en français en te basant sur le contenu VISUEL (OCR du texte, "
            "schémas, UI, code, erreurs). Sois précis et actionnable."
        )
    else:
        prompt_vision = (
            "Analyse cette image en français de façon précise : "
            "1) texte visible (OCR), 2) éléments UI / schéma / code, "
            "3) ce que montre l'image, 4) points utiles pour l'utilisateur."
        )
    for f in fichiers_traites:
        if f.get("type") != "image" or not f.get("image_b64"):
            continue
        if f.get("vision_ok") and f.get("contenu") and not str(f["contenu"]).startswith("[Image"):
            continue
        description, _tokens = client.analyser_image(
            image_b64=f["image_b64"],
            media_type=f.get("media_type") or "image/png",
            question=prompt_vision,
        )
        if description and not description.startswith("[Vision indisponible"):
            f["contenu"] = description
            f["vision_ok"] = True
        else:
            f["contenu"] = description or f.get("contenu") or f"[Image : {f.get('nom')}]"
            f["vision_ok"] = False
    return fichiers_traites


def formater_pour_prompt(fichiers_traites: list) -> str:
    """Formate les fichiers uploadés pour injection dans le prompt"""
    if not fichiers_traites:
        return ""

    blocs = []
    for f in fichiers_traites:
        if f["type"] == "image":
            if f.get("vision_ok") and f.get("contenu"):
                blocs.append(
                    f"[Image analysée : {f['nom']}]\n{f['contenu']}"
                )
            elif f.get("contenu") and not str(f["contenu"]).startswith("[Image"):
                blocs.append(
                    f"[Image : {f['nom']}]\n{f['contenu']}"
                )
            else:
                blocs.append(
                    f"[Fichier image : {f['nom']}] — analyse visuelle en attente"
                )
        elif f["contenu"]:
            lang = f["langage"]
            if f["type"] == "code":
                blocs.append(
                    f"[Fichier {f['nom']} — {lang}]\n"
                    f"```{lang}\n{f['contenu']}\n```"
                )
            elif f["type"] == "word":
                blocs.append(
                    f"[Document Word : {f['nom']}]\n{f['contenu']}"
                )
            else:
                blocs.append(
                    f"[Fichier {f['nom']} — {lang}]\n{f['contenu']}"
                )

    return "\n\n" + ("─" * 50 + "\n").join(blocs) if blocs else ""

def transcrire_audio(fichier_audio, modele: str = "whisper-large-v3") -> str:
    """Transcrit un enregistrement audio via Groq Whisper. Retourne le texte."""
    import os
    from groq import Groq

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return ""

    nom = getattr(fichier_audio, "name", "audio.wav") or "audio.wav"
    if hasattr(fichier_audio, "getvalue"):
        data = fichier_audio.getvalue()
    elif hasattr(fichier_audio, "read"):
        data = fichier_audio.read()
        if hasattr(fichier_audio, "seek"):
            try:
                fichier_audio.seek(0)
            except Exception:
                pass
    else:
        data = fichier_audio

    client = Groq(api_key=api_key)
    transcription = client.audio.transcriptions.create(
        file=(nom, data),
        model=modele,
        language="fr",
        response_format="text",
    )
    if isinstance(transcription, str):
        return transcription.strip()
    return str(getattr(transcription, "text", transcription) or "").strip()
